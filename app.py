# import os
# import re
# import smtplib
# import pandas as pd
# from flask import Flask, render_template, request, redirect, url_for, flash, send_file
# from werkzeug.utils import secure_filename
# from email.message import EmailMessage
# from PyPDF2 import PdfReader
# from docx import Document
# from datetime import datetime
# from dotenv import load_dotenv

# load_dotenv()

# app = Flask(__name__)
# app.secret_key = "supersecretkey"

# UPLOAD_FOLDER = "uploads"
# ALLOWED_EXTENSIONS = {"txt", "csv", "docx", "pdf"}
# app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

# EMAIL_SENDER = os.getenv("EMAIL_SENDER")
# EMAIL_PASSWORD = os.getenv("EMAIL_PASSWORD")

# if not os.path.exists(UPLOAD_FOLDER):
#     os.makedirs(UPLOAD_FOLDER)

# def allowed_file(filename):
#     return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

# def extract_emails_from_text(text):
#     return list(set(re.findall(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)))

# def extract_emails_from_file(filepath):
#     emails = []
#     file_extension = filepath.rsplit(".", 1)[1].lower()

#     try:
#         if file_extension == "txt":
#             with open(filepath, "r", encoding="utf-8") as f:
#                 emails = extract_emails_from_text(f.read())
#         elif file_extension == "csv":
#             df = pd.read_csv(filepath, encoding_errors='ignore')
#             for column in df.columns:
#                 emails.extend(extract_emails_from_text(" ".join(df[column].astype(str))))
#         elif file_extension == "docx":
#             doc = Document(filepath)
#             text = "\n".join([para.text for para in doc.paragraphs])
#             emails = extract_emails_from_text(text)
#         elif file_extension == "pdf":
#             pdf_reader = PdfReader(filepath)
#             text = "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
#             emails = extract_emails_from_text(text)
#     except Exception as e:
#         print(f"Error processing {filepath}: {e}")
    
#     return list(set(emails))

# def send_email(to_email, resume_path):
#     subject = "Application for Software Tester Position"
#     body = """
# Dear Hiring Manager,

# I hope you are doing well. I am writing to express my interest in the Software Tester position at your company.

# Looking forward to your response.

# Best regards,
# Chinthalapalli Basi Reddy
# basi.reddy.c1996@gmail.com
# +91 9494912885
# """
#     msg = EmailMessage()
#     msg["Subject"] = subject
#     msg["From"] = EMAIL_SENDER
#     msg["To"] = to_email
#     msg.set_content(body)
    
#     try:
#         # Read resume and attach it
#         with open(resume_path, "rb") as f:
#             resume_data = f.read()
#             # Adjust maintype and subtype if your resume is not a PDF
#             msg.add_attachment(resume_data, maintype="application", subtype="pdf", filename="Resume.pdf")
        
#         with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
#             server.login(EMAIL_SENDER, EMAIL_PASSWORD)
#             server.send_message(msg)
#         return "Sent"
#     except Exception as e:
#         print(f"Error sending email to {to_email}: {e}")
#         return "Failed"

# @app.route("/", methods=["GET", "POST"])
# def index():
#     email_status_list = []
    
#     if request.method == "POST":
#         if "file" not in request.files or "resume" not in request.files:
#             flash("Missing file or resume.")
#             return redirect(request.url)
        
#         file = request.files["file"]
#         resume = request.files["resume"]
        
#         if file.filename == "" or resume.filename == "":
#             flash("No file or resume selected.")
#             return redirect(request.url)

#         if file and allowed_file(file.filename) and resume and allowed_file(resume.filename):
#             file_path = os.path.join(app.config["UPLOAD_FOLDER"], secure_filename(file.filename))
#             resume_path = os.path.join(app.config["UPLOAD_FOLDER"], "Resume.pdf")
#             file.save(file_path)
#             resume.save(resume_path)

#             emails = extract_emails_from_file(file_path)
#             if emails:
#                 for i, email in enumerate(emails):
#                     status = send_email(email, resume_path)
#                     email_status_list.append({
#                         "sno": i + 1,
#                         "email": email,
#                         "body": "Application for Software Tester Position",
#                         "status": status,
#                         "date_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
#                     })
#                 df = pd.DataFrame(email_status_list)
#                 csv_path = os.path.join(app.config["UPLOAD_FOLDER"], "email_status.csv")
#                 df.to_csv(csv_path, index=False)
#                 flash(f"Emails processed successfully! {len(emails)} extracted.")
#             else:
#                 flash("No valid emails found in the file.")
    
#     return render_template("index.html", email_status_list=email_status_list)

# @app.route("/download")
# def download_file():
#     path = os.path.join(app.config["UPLOAD_FOLDER"], "email_status.csv")
#     if os.path.exists(path):
#         return send_file(path, as_attachment=True)
#     else:
#         flash("No email status file available.")
#         return redirect(url_for("index"))

# if __name__ == "__main__":
#     app.run(debug=True)











import os
import re
import smtplib
import pandas as pd
from flask import Flask, render_template, request, redirect, url_for, flash, send_file
from werkzeug.utils import secure_filename
from email.message import EmailMessage
from PyPDF2 import PdfReader
from docx import Document
from datetime import datetime
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)
app.secret_key = "supersecretkey"

UPLOAD_FOLDER = "uploads"
ALLOWED_EXTENSIONS = {"txt", "csv", "docx", "pdf"}
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
EMAIL_SENDER = os.getenv("EMAIL_SENDER")
EMAIL_PASSWORD = os.getenv("EMAIL_PASSWORD")


# EMAIL_SENDER = os.getenv("basi.reddy.c1996@gmail.com")
# EMAIL_PASSWORD = os.getenv("ifyl oiuf ftdl wnzd")

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_emails_from_text(text):
    return list(set(re.findall(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)))

def extract_emails_from_file(filepath):
    emails = []
    file_extension = filepath.rsplit(".", 1)[1].lower()

    try:
        if file_extension == "txt":
            with open(filepath, "r", encoding="utf-8") as f:
                emails = extract_emails_from_text(f.read())
        elif file_extension == "csv":
            df = pd.read_csv(filepath, encoding_errors='ignore')
            for column in df.columns:
                emails.extend(extract_emails_from_text(" ".join(df[column].astype(str))))
        elif file_extension == "docx":
            doc = Document(filepath)
            text = "\n".join([para.text for para in doc.paragraphs])
            emails = extract_emails_from_text(text)
        elif file_extension == "pdf":
            pdf_reader = PdfReader(filepath)
            text = "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
            emails = extract_emails_from_text(text)
    except Exception as e:
        print(f"Error processing {filepath}: {e}")
    
    return list(set(emails))

def send_email(to_email, resume_path):
    subject = "Application for Software Tester Position"
    body = """
Dear Hiring Manager,

I hope you are doing well. I am writing to express my interest in the Software Tester position at your company.

Looking forward to your response.

Best regards,
Chinthalapalli Basi Reddy
basi.reddy.c1996@gmail.com
+91 9494912885
"""
    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = EMAIL_SENDER
    msg["To"] = to_email
    msg.set_content(body)
    
    try:
        # Read resume and attach it
        with open(resume_path, "rb") as f:
            resume_data = f.read()
            # Adjust maintype/subtype if your resume is not a PDF
            msg.add_attachment(resume_data, maintype="application", subtype="pdf", filename="Resume.pdf")
        
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(EMAIL_SENDER, EMAIL_PASSWORD)
            server.send_message(msg)
        return "Sent"
    except Exception as e:
        print(f"Error sending email to {to_email}: {e}")
        return "Failed"

@app.route("/", methods=["GET", "POST"])
def index():
    email_status_list = []
    
    if request.method == "POST":
        if "file" not in request.files or "resume" not in request.files:
            flash("Missing file or resume.")
            return redirect(request.url)
        
        file = request.files["file"]
        resume = request.files["resume"]
        
        if file.filename == "" or resume.filename == "":
            flash("No file or resume selected.")
            return redirect(request.url)

        if file and allowed_file(file.filename) and resume and allowed_file(resume.filename):
            file_path = os.path.join(app.config["UPLOAD_FOLDER"], secure_filename(file.filename))
            resume_path = os.path.join(app.config["UPLOAD_FOLDER"], "Resume.pdf")
            file.save(file_path)
            resume.save(resume_path)

            emails = extract_emails_from_file(file_path)
            if emails:
                for i, email in enumerate(emails):
                    status = send_email(email, resume_path)
                    email_status_list.append({
                        "sno": i + 1,
                        "email": email,
                        "body": "Application for Software Tester Position",
                        "status": status,
                        "date_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    })
                df = pd.DataFrame(email_status_list)
                csv_path = os.path.join(app.config["UPLOAD_FOLDER"], "email_status.csv")
                df.to_csv(csv_path, index=False)
                flash(f"Emails processed successfully! {len(emails)} extracted.")
            else:
                flash("No valid emails found in the file.")
    
    return render_template("index.html", email_status_list=email_status_list)

@app.route("/download")
def download_file():
    path = os.path.join(app.config["UPLOAD_FOLDER"], "email_status.csv")
    if os.path.exists(path):
        return send_file(path, as_attachment=True)
    else:
        flash("No email status file available.")
        return redirect(url_for("index"))

if __name__ == "__main__":
    app.run(debug=True)
